import os
import ast
from docx import Document

#Input and interface
import tkinter as tk
from tkinter import filedialog

my_w = tk.Tk()
my_w.geometry("450x200")
my_w.title("Docx_stylize.exe")
my_dir=''


def my_fun():
    my_dir = filedialog.askopenfilename()
    if my_dir:
        b1.config(text=my_dir)
    my_fun.my_dir = my_dir.replace('/', '\\')

def my_fun2():
    my_dir1 = filedialog.askdirectory()
    if my_dir1:
        b2.config(text=my_dir1)
    my_fun2.my_dir1 = my_dir1.replace('/', '\\')

def data():
    #File path
    with open(my_fun.my_dir, 'r') as file:
        lines = file.readlines()
        print(lines)
    lines = lines.pop()
    print(lines)
    x = lines.split('|')
    wall_names = []
    wall_widths = []
    wall_materials = []
    wall_names = ast.literal_eval(x[0])
    wall_widths = ast.literal_eval(x[1])
    wall_materials = ast.literal_eval(x[2])
    print(wall_widths)
    print(wall_names)
    print(wall_materials)
    #Work inside document
    document = Document()
    for i in range(len((wall_names))):
        document.add_heading(wall_names[i], level=1)
        p = document.add_paragraph()
        p.add_run('Width [cm]' + '\t' + 'Material').bold = True
        for y in range(len(wall_widths[i])):
            document.add_paragraph(wall_widths[i][y] + '\t' + wall_materials[i][y])

    for i in range(len((wall_names))):
        document.add_heading(wall_names[i], level=1)
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Width [cm]'
        hdr_cells[1].text = 'Material'
        for y in range(len(wall_widths[i])):
            row_cells = table.add_row().cells
            row_cells[0].text = wall_widths[i][y]
            row_cells[1].text = wall_materials[i][y]
    document.save(my_fun2.my_dir1 + '\\' + 'walltype.docx')
    l2 = tk.Label(my_w, text='Data successfully converted.', width=40, font=my_font1)
    l2.grid(row=4,column=0)

#Window and button setup
my_font1 = ('arial', 10, 'normal')
l1 = tk.Label(my_w, text='Stylizing the wall type export into a docx file', width=40, font=my_font1)
l1.grid(row=0,column=0, padx=20, pady=20)
b1 = tk.Button(my_w, text='Select the source file', width=55, command=lambda:my_fun())
b1.grid(row=1, column=0, padx=20)
b2 = tk.Button(my_w, text='Select directory to save to', width=55, command=lambda:my_fun2())
b2.grid(row=2, column=0, padx=20)
b3 = tk.Button(my_w, text='Create the document', width=55, command=lambda:data())
b3.grid(row=4, column=0, padx=20, pady=20)

#Displaying the window
my_w.mainloop()



